import io
import logging
from typing import BinaryIO

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text_parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_text(file: BinaryIO, filename: str) -> tuple[str, str]:
    file_type = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    file_bytes = file.read()
    if file_type == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        text = extract_text_from_docx(file_bytes)
    elif file_type == "doc":
        raise ValueError("DOC格式暂不支持，请转换为DOCX或PDF后上传")
    else:
        raise ValueError(f"不支持的文件格式: {file_type}，仅支持PDF和DOCX")
    if not text.strip():
        raise ValueError("无法从文件中提取文本内容，请检查文件是否损坏或为扫描件")
    return text, file_type
